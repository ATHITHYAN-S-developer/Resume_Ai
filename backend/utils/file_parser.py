"""
File Parser Utility
Extracts text from PDF and DOCX resume files
"""

import io
import pdfplumber
import docx


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from uploaded PDF or DOCX file.
    Returns extracted text string.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload PDF or DOCX.")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")
